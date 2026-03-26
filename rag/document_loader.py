"""Document loading utilities for PDF and DOCX files."""

from langchain.schema import Document
from pypdf import PdfReader
from docx import Document as DocxDocument


def load_document(uploaded_file) -> list[Document]:
    """
    Load a document from a Streamlit UploadedFile object.

    Supports PDF and DOCX formats. Returns a list of LangChain Document
    objects with metadata including filename and page number.
    """
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return _load_pdf(uploaded_file)
    elif filename.endswith(".docx"):
        return _load_docx(uploaded_file)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Use PDF or DOCX.")


def _load_pdf(uploaded_file) -> list[Document]:
    """Extract text from a PDF file page by page."""
    reader = PdfReader(uploaded_file)
    documents = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": uploaded_file.name,
                        "page": i + 1,
                        "total_pages": len(reader.pages),
                    },
                )
            )

    return documents


def _load_docx(uploaded_file) -> list[Document]:
    """Extract text from a DOCX file."""
    doc = DocxDocument(uploaded_file)
    full_text = "\n".join(
        paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
    )

    if full_text.strip():
        return [
            Document(
                page_content=full_text,
                metadata={"source": uploaded_file.name, "page": 1},
            )
        ]

    return []
